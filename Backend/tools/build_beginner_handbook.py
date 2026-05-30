from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from pathlib import Path


OUT = Path("Online_Grocery_Store_Backend_Beginner_Handbook.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 28, 38)
MUTED = RGBColor(90, 96, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Pt(widths[i] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Online Grocery Store Backend Handbook")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 24, INK, 0, 10),
        ("Subtitle", 13, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        if style_name.startswith("Heading"):
            style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(9.5)
    code.font.color.rgb = RGBColor(45, 45, 45)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.left_indent = Inches(0.15)

    callout = styles.add_style("Callout", 1)
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = INK
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.left_indent = Inches(0.15)
    callout.paragraph_format.right_indent = Inches(0.15)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text="", style=None):
    return doc.add_paragraph(text, style=style)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc, text):
    for line in text.strip("\n").splitlines():
        p(doc, line, "CodeBlock")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = head
        set_cell_shading(cell, LIGHT_BLUE)
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
    set_table_width(t, widths)
    p(doc, "")
    return t


def callout(doc, title, body):
    para = p(doc, style="Callout")
    run = para.add_run(title + ": ")
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    para.add_run(body)


def cover(doc):
    p(doc, "Project-Based Backend Development", None).alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Online Grocery Store Backend Handbook")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = INK
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("A beginner-friendly guide to Spring Boot, MySQL, JWT, OAuth2, Redis, WebSocket, Swagger, and backend architecture")
    sr.font.size = Pt(12)
    sr.font.color.rgb = MUTED
    p(doc, "")
    table(doc, ["Project", "Audience", "Goal"], [[
        "Online Grocery Store Platform",
        "Learners with zero or little backend experience",
        "Understand, build, explain, and interview confidently"
    ]], [2500, 3300, 3560])
    callout(doc, "How to read this handbook", "Do not try to memorize everything in one pass. First understand the story: request comes in, controller receives it, service applies rules, repository talks to database, security protects the path.")
    doc.add_page_break()


def toc(doc):
    h(doc, "Table of Contents", 1)
    sections = [
        "1. Backend Development From Zero",
        "2. The Online Grocery Store Project",
        "3. Spring Boot Project Setup",
        "4. Layered Architecture",
        "5. Entities and Database Design",
        "6. Repository Layer",
        "7. DTOs and Validation",
        "8. Service Layer and Business Logic",
        "9. Controller Layer and REST APIs",
        "10. Global Exception Handling",
        "11. Security With JWT",
        "12. OAuth2 Login",
        "13. Cart and Order Flow",
        "14. Redis Caching",
        "15. WebSocket Real-Time Updates",
        "16. Swagger/OpenAPI",
        "17. Testing",
        "18. Interview Preparation",
        "19. How to Run the Project",
        "20. Glossary",
    ]
    nums(doc, sections)
    doc.add_page_break()


def body(doc):
    h(doc, "1. Backend Development From Zero", 1)
    p(doc, "Backend development is the part of software engineering that handles data, rules, security, and communication between the frontend and the database. If a user clicks Add to Cart on a grocery website, the backend decides whether the product exists, whether stock is available, how the cart should change, and what response should go back.")
    table(doc, ["Frontend", "Backend", "Database"], [[
        "Screens, buttons, forms, product cards",
        "APIs, business rules, security, validation",
        "Permanent storage for users, products, carts, orders"
    ]], [2800, 3300, 3260])
    callout(doc, "Real-life example", "A frontend is like the grocery store display area. The backend is the billing counter, inventory manager, and rule book. The database is the storage room and accounting register.")

    h(doc, "What happens when a user places an order?", 2)
    nums(doc, [
        "The frontend sends an HTTP request to the backend.",
        "The controller receives the request.",
        "The service checks business rules such as stock and empty cart.",
        "The repository reads or writes database records.",
        "The backend returns a response such as order placed successfully.",
    ])
    code(doc, """
User action: Click Place Order
HTTP request: POST /api/orders
Backend response: 201 Created + order details
""")

    h(doc, "2. The Online Grocery Store Project", 1)
    p(doc, "This project is an online grocery backend. It allows users to register, log in, browse products, add items to cart, place orders, and receive real-time order status updates.")
    table(doc, ["Feature", "What it means in the project"], [
        ["Authentication", "Users register and log in with JWT tokens."],
        ["Catalog", "Admin creates categories and products."],
        ["Cart", "Customer adds products and quantity."],
        ["Order", "Customer converts cart into an order."],
        ["Caching", "Redis can speed up product/category reads."],
        ["Real-time updates", "WebSocket notifies clients when order status changes."],
    ], [2500, 6860])

    h(doc, "3. Spring Boot Project Setup", 1)
    p(doc, "Spring Boot is a framework that makes Java backend development faster. It auto-configures many common things such as embedded server, dependency wiring, JSON conversion, and database integration.")
    table(doc, ["Dependency", "Why we use it"], [
        ["spring-boot-starter-web", "Build REST APIs and run embedded Tomcat."],
        ["spring-boot-starter-data-jpa", "Work with databases using entities and repositories."],
        ["mysql-connector-j", "Connect Java application to MySQL."],
        ["spring-boot-starter-security", "Authentication and authorization."],
        ["jjwt", "Generate and validate JWT tokens."],
        ["spring-boot-starter-data-redis", "Use Redis for caching."],
        ["spring-boot-starter-websocket", "Real-time communication with clients."],
        ["lombok", "Reduce boilerplate getters, setters, constructors."],
    ], [3000, 6360])
    code(doc, """
src/main/java/com/grocery/store
  controller
  service
  repository
  entity
  dto
  security
  config
  exception
""")
    callout(doc, "Interview explanation", "I set up a layered Spring Boot project with dependencies for web APIs, JPA, MySQL, validation, security, JWT, Redis, WebSocket, Swagger, and tests. This keeps the project modular and easy to extend.")

    h(doc, "4. Layered Architecture", 1)
    p(doc, "Layered architecture means each part of the backend has a specific responsibility. This avoids mixing API code, business rules, and database code in one place.")
    code(doc, """
HTTP Request
    -> Controller
        -> Service
            -> Repository
                -> Entity
                    -> MySQL
""")
    table(doc, ["Layer", "Responsibility", "Example"], [
        ["Controller", "Receives HTTP requests", "CartController.addProduct"],
        ["Service", "Business rules", "CartService checks stock"],
        ["Repository", "Database access", "ProductRepository.findByActiveTrue"],
        ["Entity", "Database table model", "Product, User, Order"],
        ["DTO", "API request/response shape", "RegisterRequest, UserResponse"],
    ], [1700, 3300, 4360])
    callout(doc, "Why it matters", "If all logic is in controllers, the code becomes hard to test and hard to explain. Services make business logic reusable and repositories keep database access clean.")

    h(doc, "5. Entities and Database Design", 1)
    p(doc, "An entity is a Java class mapped to a database table. In this project, User maps to the users table, Product maps to products, and Order maps to orders.")
    table(doc, ["Entity", "Purpose"], [
        ["User", "Stores customer/admin account details."],
        ["Category", "Groups products such as Fruits or Dairy."],
        ["Product", "Stores grocery item name, price, stock, active status."],
        ["Cart", "One active cart for a user."],
        ["CartItem", "Product and quantity inside a cart."],
        ["Order", "Completed purchase summary."],
        ["OrderItem", "Purchased product snapshot with price at purchase."],
    ], [2200, 7160])
    h(doc, "Important JPA annotations", 2)
    table(doc, ["Annotation", "Meaning"], [
        ["@Entity", "This class maps to a database table."],
        ["@Id", "Primary key column."],
        ["@GeneratedValue", "Database generates the id."],
        ["@Column", "Column details such as nullable or unique."],
        ["@ManyToOne", "Many records point to one parent, such as many products in one category."],
        ["@OneToOne", "One record maps to one other record, such as one cart per user."],
        ["@Enumerated", "Stores enum values such as CUSTOMER or ADMIN."],
    ], [2300, 7060])
    callout(doc, "Why priceAtPurchase exists", "Product price can change later. Old orders must keep the price that was paid at that time, so OrderItem stores priceAtPurchase.")

    h(doc, "6. Repository Layer", 1)
    p(doc, "Repositories are interfaces that Spring Data JPA uses to generate database queries. You do not need to write SQL for common CRUD operations.")
    code(doc, """
public interface ProductRepository extends JpaRepository<Product, Long> {
    List<Product> findByActiveTrue();
    List<Product> findByCategoryIdAndActiveTrue(Long categoryId);
    List<Product> findByNameContainingIgnoreCaseAndActiveTrue(String name);
}
""")
    p(doc, "Spring reads method names and creates queries automatically. For example, findByEmail means select user where email equals the given value.")
    callout(doc, "Interview explanation", "Repositories separate database access from business logic. They make common CRUD operations available through JpaRepository and support custom derived queries.")

    h(doc, "7. DTOs and Validation", 1)
    p(doc, "DTO means Data Transfer Object. DTOs define the shape of data entering and leaving the API. They protect entities from being exposed directly.")
    table(doc, ["DTO Type", "Example", "Purpose"], [
        ["Request DTO", "RegisterRequest", "Data sent by client to backend."],
        ["Response DTO", "UserResponse", "Data returned by backend to client."],
        ["Error DTO", "ErrorResponse", "Consistent error format."],
    ], [2000, 2500, 4860])
    code(doc, """
public record RegisterRequest(
    @NotBlank String fullName,
    @Email @NotBlank String email,
    @Size(min = 6) String password,
    @NotBlank String phoneNumber
) {}
""")
    callout(doc, "Why not return User entity directly", "User contains password. Even if it is hashed, it should never be returned in API responses. UserResponse exposes only safe fields.")

    h(doc, "8. Service Layer and Business Logic", 1)
    p(doc, "Services contain rules. A controller should not know how to calculate totals or reduce stock. That belongs in services.")
    table(doc, ["Service", "Rules handled"], [
        ["UserService", "Register customer/admin, hash password, find user."],
        ["CategoryService", "Create category, prevent duplicate names, cache reads."],
        ["ProductService", "Create product, search products, reduce stock."],
        ["CartService", "Create cart, add item, validate stock, clear cart."],
        ["OrderService", "Place order, calculate total, create order items, update status."],
    ], [2300, 7060])
    h(doc, "Why @Transactional matters", 2)
    p(doc, "Order placement is multiple database operations. If one operation fails after some records were saved, the database could become inconsistent. @Transactional makes the entire method succeed or roll back together.")
    code(doc, """
placeOrder()
  get cart items
  calculate total
  save order
  save order items
  reduce product stock
  clear cart
""")

    h(doc, "9. Controller Layer and REST APIs", 1)
    p(doc, "Controllers expose endpoints. They receive requests, validate DTOs, call services, and return response DTOs.")
    code(doc, """
@PostMapping("/api/cart/items")
public CartItemResponse addProduct(@Valid @RequestBody AddToCartRequest request) {
    CartItem cartItem = cartService.addProduct(
        request.userId(),
        request.productId(),
        request.quantity()
    );
    return CartItemResponse.from(cartItem);
}
""")
    table(doc, ["API", "Purpose", "Access"], [
        ["POST /api/auth/register", "Register customer", "Public"],
        ["POST /api/auth/login", "Login and receive tokens", "Public"],
        ["POST /api/categories", "Create category", "Admin"],
        ["POST /api/products", "Create product", "Admin"],
        ["POST /api/cart/items", "Add item to cart", "Customer/Admin"],
        ["POST /api/orders", "Place order", "Customer/Admin"],
        ["PATCH /api/orders/{id}/status", "Update order status", "Admin"],
    ], [3000, 4100, 2260])

    h(doc, "10. Global Exception Handling", 1)
    p(doc, "Global exception handling means errors are handled in one place using @RestControllerAdvice. This avoids writing try-catch blocks in every controller.")
    table(doc, ["Exception", "HTTP Status", "Meaning"], [
        ["ResourceNotFoundException", "404", "Record does not exist."],
        ["DuplicateResourceException", "409", "Email or category already exists."],
        ["BusinessRuleException", "400", "Valid request shape but rule failed."],
        ["MethodArgumentNotValidException", "400", "DTO validation failed."],
        ["AuthenticationException", "401", "Invalid login or token."],
    ], [3200, 1800, 4360])
    callout(doc, "Example", "If the cart is empty and the user tries to place an order, the API returns a clean 400 response saying Cart is empty.")

    h(doc, "11. Security With JWT", 1)
    p(doc, "JWT means JSON Web Token. After login, the backend gives the client a token. The client sends that token with future requests.")
    code(doc, """
Authorization: Bearer <token>
""")
    nums(doc, [
        "User logs in with email and password.",
        "Spring Security verifies credentials.",
        "Backend generates JWT.",
        "Client stores JWT.",
        "Client sends JWT in Authorization header.",
        "JwtAuthenticationFilter validates token and sets SecurityContext.",
    ])
    h(doc, "Stateless security", 2)
    p(doc, "Stateless means the server does not store login session data for every user. The token carries enough information to identify the user for each request.")
    table(doc, ["Concept", "Meaning"], [
        ["Access token", "Shorter-lived token used for API calls."],
        ["Refresh token", "Longer-lived token used to get a new access token."],
        ["BCrypt", "Password hashing algorithm. Stores hash, not plain password."],
        ["SecurityContext", "Spring Security holder for current authenticated user."],
    ], [2500, 6860])

    h(doc, "12. OAuth2 Login", 1)
    p(doc, "OAuth2 lets users sign in using a trusted provider such as Google. Google proves the user's identity, but our backend still issues its own JWT for grocery API access.")
    code(doc, """
/oauth2/authorization/google
        -> Google login
        -> /login/oauth2/code/google
        -> create or fetch app user
        -> return app JWT
""")
    callout(doc, "Why create our own JWT", "Google login proves who the user is. Our JWT controls what the user can do inside our grocery platform and keeps the rest of the backend stateless.")

    h(doc, "13. Cart and Order Flow", 1)
    p(doc, "Cart and order flow is the heart of the grocery platform. It is where entity relationships, services, repositories, transactions, and security all work together.")
    code(doc, """
Customer adds Apple quantity 2
  -> validate product exists
  -> validate stock is enough
  -> create or update cart item

Customer places order
  -> validate cart is not empty
  -> calculate total
  -> create order
  -> create order items with priceAtPurchase
  -> reduce stock
  -> clear cart
""")
    table(doc, ["Rule", "Why"], [
        ["Cannot order empty cart", "Prevents meaningless orders."],
        ["Cannot exceed stock", "Prevents selling unavailable products."],
        ["Reduce stock after order", "Inventory must stay accurate."],
        ["Clear cart after order", "Cart should not contain already purchased items."],
        ["Store priceAtPurchase", "Old order history remains correct."],
    ], [3000, 6360])

    h(doc, "14. Redis Caching", 1)
    p(doc, "Redis is an in-memory data store. It is often used to cache frequently read data so the application does not repeatedly query the database for the same information.")
    table(doc, ["Without Cache", "With Redis Cache"], [[
        "Every product list request hits MySQL.",
        "First request hits MySQL; later repeated requests can be served from cache."
    ]], [4680, 4680])
    code(doc, """
Redis local address:
localhost:6387

Run with Redis profile:
mvn spring-boot:run -Dspring-boot.run.profiles=redis
""")
    p(doc, "In this project, product and category reads are cacheable. Cache is evicted when products/categories change or stock changes after order placement.")
    callout(doc, "Interview explanation", "I used Spring Cache abstraction so the service code is not tightly coupled to Redis. The default profile uses simple cache for easy local startup, and the redis profile switches to Redis.")

    h(doc, "15. WebSocket Real-Time Updates", 1)
    p(doc, "REST APIs are request-response. WebSocket keeps a connection open so the server can push updates to the client. This is useful for order status updates.")
    code(doc, """
WebSocket endpoint: /ws
Subscribe topic:   /topic/orders/{orderId}

Admin updates order status:
PATCH /api/orders/{orderId}/status

Backend publishes:
{
  "orderId": 1,
  "status": "CONFIRMED",
  "updatedAt": "..."
}
""")
    table(doc, ["REST", "WebSocket"], [[
        "Client asks server for data.",
        "Server can push updates to subscribed clients."
    ]], [4680, 4680])

    h(doc, "16. Swagger/OpenAPI", 1)
    p(doc, "Swagger creates interactive API documentation. It helps developers see endpoints, request bodies, response formats, and test APIs from a browser.")
    code(doc, """
Swagger UI:
http://localhost:8080/swagger-ui.html
""")
    callout(doc, "Why it helps", "Swagger is useful for demos, team collaboration, and interview walkthroughs because the API contract is visible.")

    h(doc, "17. Testing", 1)
    p(doc, "Testing proves that the backend works. This project uses integration tests with MockMvc and H2. H2 is an in-memory database used only for tests.")
    table(doc, ["Tool", "Purpose"], [
        ["JUnit", "Runs test methods."],
        ["MockMvc", "Tests controllers without manually opening a browser."],
        ["H2", "Temporary database for tests."],
        ["SpringBootTest", "Loads the Spring application context."],
    ], [2200, 7160])
    p(doc, "The integration test checks admin registration, customer registration, login, refresh token, role protection, product creation, cart, order, stock reduction, and order status update.")
    code(doc, """
mvn test

Expected result:
BUILD SUCCESS
Tests run: 1, Failures: 0, Errors: 0
""")

    h(doc, "18. Interview Preparation", 1)
    table(doc, ["Question", "Strong answer"], [
        ["Why layered architecture?", "It separates API, business logic, and database logic, making code maintainable and testable."],
        ["Why DTOs?", "DTOs protect entities, define API contracts, and support validation."],
        ["Why service layer?", "Business rules belong in services, not controllers."],
        ["Why JWT?", "JWT enables stateless authentication for protected APIs."],
        ["Why refresh token?", "It lets users get new access tokens without logging in again."],
        ["Why Redis?", "It reduces repeated database reads for frequently accessed data."],
        ["Why WebSocket?", "It supports real-time order status updates."],
        ["Why @Transactional?", "Multi-step database operations either fully succeed or roll back."],
    ], [3000, 6360])
    h(doc, "How to explain the full project in 60 seconds", 2)
    p(doc, "I built an online grocery backend using Spring Boot and MySQL. The project follows layered architecture with controllers, services, repositories, entities, DTOs, security, and exception handling. Users can register and log in using JWT, and OAuth2 is available for Google login. Admins manage products and categories, while customers manage carts and orders. Order placement validates stock, creates order records, stores price at purchase, reduces inventory, and clears the cart in a transaction. Redis caching improves product and category reads, WebSocket pushes order status updates, Swagger documents APIs, and integration tests verify the main flow.")

    h(doc, "19. How to Run the Project", 1)
    code(doc, """
Create database:
CREATE DATABASE grocery_store;

Run application:
mvn spring-boot:run

Run with Redis:
mvn spring-boot:run -Dspring-boot.run.profiles=redis

Run tests:
mvn test
""")
    table(doc, ["Profile", "Purpose"], [
        ["default", "MySQL + simple in-memory cache."],
        ["redis", "MySQL + Redis cache on localhost:6387."],
        ["oauth2", "Google OAuth2 login with credentials."],
        ["test", "H2 database for integration tests."],
    ], [2200, 7160])

    h(doc, "20. Glossary", 1)
    table(doc, ["Term", "Simple meaning"], [
        ["API", "A way for software systems to communicate."],
        ["REST", "A style for APIs using HTTP methods like GET and POST."],
        ["HTTP", "Protocol used by web clients and servers."],
        ["JSON", "Text format commonly used in API requests/responses."],
        ["Entity", "Java class mapped to a database table."],
        ["Repository", "Database access interface."],
        ["Service", "Business logic class."],
        ["Controller", "API entry point."],
        ["DTO", "Object used for request/response data."],
        ["JWT", "Token used for stateless authentication."],
        ["OAuth2", "Login flow using providers like Google."],
        ["Redis", "Fast in-memory cache/data store."],
        ["WebSocket", "Long-lived connection for real-time updates."],
        ["Swagger", "Interactive API documentation."],
    ], [2000, 7360])


def build():
    doc = Document()
    style_doc(doc)
    cover(doc)
    toc(doc)
    body(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build()
